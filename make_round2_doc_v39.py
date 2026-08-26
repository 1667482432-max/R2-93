from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_OUT = Path(r"D:\Phase93_work\docs\Round2_Phase93_端到端模型算法说明.docx")
DESKTOP_OUT = Path(r"C:\Users\asus\Desktop\Round2_完整版.docx")

BLUE = "1976E9"
GREEN = "16895F"
ORANGE = "DB9410"
RED = "C83A3A"
DARK = "111820"
MUTED = "667085"
LINE = "CDD5DF"
PANEL = "F2F4F7"
PALE = "EAF2FB"
WARN = "FFF2D6"


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    flag = OxmlElement("w:tblHeader")
    flag.set(qn("w:val"), "true")
    tr_pr.append(flag)


def prevent_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    flag = OxmlElement("w:cantSplit")
    flag.set(qn("w:val"), "true")
    tr_pr.append(flag)


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=85, start=110, bottom=85, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_inches, indent_dxa=110):
    widths = [int(round(v * 1440)) for v in widths_inches]
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_split(row)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)


def cell_text(cell, text, bold=False, color=DARK, size=9.2):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.02
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_repeat_header(table.rows[0])
    for i, value in enumerate(headers):
        cell_text(table.rows[0].cells[i], value, True, DARK, 9.2)
        shade(table.rows[0].cells[i], PALE)
    for ridx, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cell_text(cells[i], value, False, DARK, 8.9)
            if ridx % 2:
                shade(cells[i], "FAFBFC")
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def formula(doc, value):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(value)
    r.font.name = "Cambria Math"
    r.font.size = Pt(11.5)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK)


def callout(doc, label, body, fill=PALE, label_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    a = p.add_run(label + "：")
    a.bold = True
    a.font.color.rgb = RGBColor.from_string(label_color)
    b = p.add_run(body)
    b.font.color.rgb = RGBColor.from_string(DARK)
    for r in p.runs:
        r.font.name = "Microsoft YaHei"
        r.font.size = Pt(10)
    set_table_geometry(table, [6.95])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def _new_numbering_id(doc):
    root = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in root.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in root.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    lvl.append(fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "360")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "180")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    first_num_index = next(
        (i for i, child in enumerate(list(root)) if child.tag == qn("w:num")),
        len(root),
    )
    root.insert(first_num_index, abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    root.append(num)
    return num_id


def numbered_list(doc, items):
    num_id = _new_numbering_id(doc)
    for text in items:
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num])
        p_pr.append(num_pr)
        p.add_run(text)


def paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def chapter(doc, number, title, subtitle):
    doc.add_heading(f"CHAPTER {number:02d}", level=1)
    p = doc.add_paragraph()
    p.style = doc.styles["Title 2"]
    p.add_run(title)
    q = doc.add_paragraph(subtitle)
    q.style = doc.styles["Subtitle"]


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.72)
sec.bottom_margin = Inches(0.68)
sec.left_margin = Inches(0.78)
sec.right_margin = Inches(0.78)

styles = doc.styles
styles["Normal"].font.name = "Microsoft YaHei"
styles["Normal"].font.size = Pt(10.2)
styles["Normal"].paragraph_format.space_after = Pt(5.5)
styles["Normal"].paragraph_format.line_spacing = 1.12
for name, size, color, before, after in [
    ("Heading 1", 10.5, BLUE, 18, 3),
    ("Heading 2", 14.5, DARK, 10, 6),
    ("Heading 3", 11.5, BLUE, 8, 4),
]:
    st = styles[name]
    st.font.name = "Microsoft YaHei"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
try:
    styles["Title 2"]
except KeyError:
    styles.add_style("Title 2", 1)
styles["Title 2"].font.name = "Microsoft YaHei"
styles["Title 2"].font.size = Pt(20)
styles["Title 2"].font.bold = True
styles["Title 2"].font.color.rgb = RGBColor.from_string(DARK)
styles["Subtitle"].font.name = "Microsoft YaHei"
styles["Subtitle"].font.size = Pt(9.5)
styles["Subtitle"].font.color.rgb = RGBColor.from_string(MUTED)

header_p = sec.header.paragraphs[0]
header_p.text = "Round2 · Phase93 端到端模型算法说明"
header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for run in header_p.runs:
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
footer_p = sec.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_p.add_run("Phase93  |  ")
r.font.size = Pt(8)
r.font.color.rgb = RGBColor.from_string(MUTED)
add_page_field(footer_p)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(52)
r = p.add_run("ROUND2  |  PHASE93")
r.bold = True
r.font.name = "Microsoft YaHei"
r.font.size = Pt(12)
r.font.color.rgb = RGBColor.from_string(BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(10)
r = p.add_run("Physical AI 无线数字孪生信道生成")
r.bold = True
r.font.name = "Microsoft YaHei"
r.font.size = Pt(28)
r.font.color.rgb = RGBColor.from_string(DARK)
p = doc.add_paragraph()
r = p.add_run("Phase93 端到端模型算法说明与复现手册")
r.font.name = "Microsoft YaHei"
r.font.size = Pt(17)
r.font.color.rgb = RGBColor.from_string(BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
r = p.add_run("重点说明：PAS / PDP / NMSE、矩形 OOF、分组空间预测、P9/P10/Phase40、g5/g6 full-192 修正、原始数据端到端复现")
r.font.name = "Microsoft YaHei"
r.font.size = Pt(11)
r.font.color.rgb = RGBColor.from_string(MUTED)
callout(doc, "阅读目标", "读完后能够从一个测试点出发，解释它为什么属于某个几何组、如何选择 PAS/PDP 邻居、怎样由 Phase6 形成 P9/P10、为什么 Phase40 反向于 P10，以及 Phase93 为什么只改变 89 个测试点。", PALE, BLUE)
callout(doc, "严格边界", "项目只读取 Round2_Train_Channel.npy、Round2_Train_Pos.npy、Round2_Test_Pos.npy、Round2_Map.ply 四个赛事原始文件。所有缓存、中间谱、OOF 折和最终 NPY 均由代码生成。", PANEL, GREEN)
doc.add_page_break()

# Chapter 1: metrics first as user requested.
chapter(doc, 1, "先理解 PAS、PDP 与 NMSE", "模型的结构由评价指标直接决定。")
paragraph(doc, "原始复信道记为 H[a,u,f]：a 是 256 个基站端口，u 是 4 个 UE 分支，f 是 192 个子载波。每个元素为复数，模表示幅度，辐角表示相位。赛事不只比较 H 的逐元素误差，还分别比较角度域和时延域的功率形状。")
add_table(doc, ["指标", "物理问题", "变换", "评价方向"], [
    ["PAS", "能量主要从哪些阵列角度到达？", "256 端口重排为 2×16×8，沿 16×8 做 FFT2", "余弦相似度越接近 1 越好"],
    ["PDP", "多径能量主要落在哪些时延？", "沿 192 子载波做正交 FFT", "余弦相似度越接近 1 越好"],
    ["NMSE", "预测复信道逐元素相差多少？", "直接比较 H_pred 与 H_true", "越接近 0 越好"],
], [0.7, 2.0, 2.75, 1.5])

doc.add_heading("1.1 PAS：阵列角度功率谱", level=2)
paragraph(doc, "PAS（Power Angular Spectrum）描述阵列接收到的功率在空间波束/角度 bin 上如何分布。本题的 256 个基站端口严格按 2 极化×16 水平阵元×8 垂直阵元解释；若排列顺序错误，PAS 与后续投影会同时失效。")
formula(doc, "A = FFT2_(H,V)(H),      P_PAS = |A|²")
formula(doc, "C_PAS = mean[ <P̂,P> / (||P̂||₂ ||P||₂) ]")
callout(doc, "简化例子", "4 阵元信道 [1,1,1,1] 的空间 FFT 能量集中在第 0 个角度 bin，归一化 PAS 可写成 [1,0,0,0]；[1,−1,1,−1] 的能量移到另一个角度 bin。两者总能量相同，但角度方向完全不同。", PALE, BLUE)
paragraph(doc, "PAS 余弦相似度只比较形状。若预测 PAS 是真值的 3 倍，归一化后方向仍相同，C_PAS 仍为 1，因此 PAS 不约束绝对幅度。")

doc.add_heading("1.2 PDP：频率/时延功率谱", level=2)
paragraph(doc, "PDP（Power Delay Profile）描述多径能量在离散时延 bin 上的分布。192 个子载波处于频率域；沿子载波轴做正交 FFT 后，频率变化被转化为时延结构。")
formula(doc, "D = FFT_frequency(H),      P_PDP = |D|²")
formula(doc, "C_PDP = mean[ <D̂_power,D_power> / (||D̂_power||₂ ||D_power||₂) ]")
callout(doc, "简化例子", "4 子载波信道 [1,1,1,1] 的时延功率集中在第 0 个 bin；[1,−1,1,−1] 集中到另一个 bin。预测与真值落在同一 bin 时相似度为 1，落在正交 bin 时接近 0。", PALE, GREEN)

doc.add_heading("1.3 NMSE：原始复信道归一化均方误差", level=2)
formula(doc, "NMSE = Σ |H_pred − H_true|² / Σ |H_true|²")
callout(doc, "数值例子", "真值 [1+j,1]，预测 [0.8+1.1j,0.9]。误差能量为 0.06，真值能量为 3，因此 NMSE=0.02。若 H_pred=−H_true，PAS/PDP 完全相同，但 NMSE 很大，说明谱正确不等于复信道正确。", WARN, RED)

doc.add_heading("1.4 综合分数", level=2)
formula(doc, "Score = 0.4·C_PAS + 0.4·C_PDP + 0.2/(1+NMSE)")
paragraph(doc, "例如 C_PAS=0.80、C_PDP=0.70、NMSE=1.00，则 Score=0.32+0.28+0.10=0.70。PAS 与 PDP 各占 40%，因此主干容量集中在谱形状；NMSE 项约束幅度与相位不要失控。")
callout(doc, "模型设计结论", "Phase93 不直接回归 196,608 个复数，而是先预测与评分一致的 PAS/PDP，再用保相位幅度投影恢复完整复信道。", PALE, BLUE)

chapter(doc, 2, "执行摘要：Round1 主线，Round2 场景化升级", "相似的是物理谱框架，变化的是验证几何与局部决策。")
add_table(doc, ["维度", "Round1 v39/v50", "Round2 Phase93", "为何要改"], [
    ["数据", "2000 训练、500 测试；16 个零信道", "4000 训练、500 测试；262 个零信道", "异常比例和有效邻域发生变化"],
    ["验证", "固定五折，每折 200 点", "5 折矩形挖空，按官方人口加权", "随机点验证高估矩形空洞泛化"],
    ["空间预测", "15 个克里金专家，按 PAS/PDP 组门控", "多个矩形区域被细分为 11 个几何组，PAS/PDP 独立邻域", "不同矩形内部、边缘与基站侧尺度不同"],
    ["地图", "209 维走廊、环形、天际线描述", "LOS/local/rich 特征进入距离、描述器与门控", "围绕 Round2 岛内验证做任务定制"],
    ["重建", "复相位初值 + 角域/时延域交替投影", "保留交替投影，采用保守幅度规范", "共同的 Physical AI 主干"],
    ["最终校准", "none/PAS/PDP/PAS+PDP 四动作", "P9 锚点残差、anti-P10、g5/g6 full-192", "只扩大有 OOF 证据的局部修正"],
], [0.7, 2.05, 2.65, 1.55])
callout(doc, "一句话", "Round1 解决“怎样从可解释谱恢复复信道”；Round2 进一步解决“官方矩形几何中，哪些点该用哪套局部规律、修正多少”。", PALE, GREEN)

chapter(doc, 3, "任务、原始文件与输出契约", "端到端复现从四个赛事文件开始。")
add_table(doc, ["文件", "作用", "是否由赛事提供"], [
    ["Round2_Train_Channel.npy", "4000 个训练点的 256×4×192 complex 信道", "是"],
    ["Round2_Train_Pos.npy", "训练位置坐标", "是"],
    ["Round2_Test_Pos.npy", "500 个测试位置坐标", "是"],
    ["Round2_Map.ply", "场景点云与法向", "是"],
    ["train_energy.npy / 各种 descriptor.npy", "训练/推理缓存", "否，由代码生成"],
    ["Phase93 NPY", "最终 (500,256,4,192) complex64", "否，由代码生成"],
], [2.4, 3.6, 1.0])
paragraph(doc, "严格入口 build_phase93_end_to_end.py 调用 build_phase93_from_raw.py。历史 P9、P10、Phase40、Phase6、PAS、PDP 或地图特征数组不能作为入口输入；Release 下载脚本仅保留用于审计历史冻结结果。")

chapter(doc, 4, "数据清洗：262 个全零信道统一剔除", "任何空间模型都不能把全零异常当成正常传播样本。")
formula(doc, "E_i = Σ_(a,u,f) |H_i[a,u,f]|²")
paragraph(doc, "若 E_i=0，则该训练点完整信道全零。Phase93 将其视为 outlier，不参与邻域树、模型拟合、矩形 OOF、锚点、校准或能量统计。共有 262 个异常点，保留 3738 个有效训练点。")
callout(doc, "一致性原则", "同一异常索引必须从所有阶段同时剔除。只在某一个训练脚本中删除，会导致邻域索引、OOF 行号、锚点数量和最终 manifest 不一致。", WARN, RED)
doc.add_page_break()
add_table(doc, ["阶段", "若保留全零点的后果"], [
    ["空间邻域", "测试点可能选择到非物理零参考，局部谱与能量被拉低"],
    ["锚点残差", "log(真值/基线) 出现极端值，残差场失真"],
    ["矩形验证", "验证人口和官方加权不再可比"],
    ["最终 QA", "可能出现低能量甚至全零输出行"],
], [1.55, 5.45])

chapter(doc, 5, "矩形 OOF：为什么不能随机划分", "验证分布必须复刻官方测试几何。")
paragraph(doc, "官方 500 个测试点集中在多个小矩形，而不是从训练区域随机抽取。随机验证点周围通常仍有非常近的训练邻居，得到的分数是“插值密集区”能力，不能代表“矩形空洞”能力。")
numbered_list(doc, [
    "从测试位置使用 DBSCAN（eps=10、min_samples=3）恢复几何分组；最终得到标签 0–10 共 11 个组。注意：代码中的 TEST_BLOCKS 保存 8 个参考矩形宽高/人口，几何分组与参考矩形数量不是同一个概念。",
    "在训练点两侧分别寻找矩形候选中心；宽、高和目标人口参考 TEST_BLOCKS，并乘 scale=0.75。",
    "验证矩形只能包含非零训练点；不同矩形不重叠，外围增加 3 m buffer，制造真实空间缺口。",
    "固定种子 20260813+fold 构造 5 折，并按官方各几何组人口对折内样本加权。",
])
callout(doc, "选型纪律", "训练过程可以复杂，但任何超参数或模块是否保留，必须依据无泄漏矩形 OOF；本地随机分数更高不能作为最终选型证据。", PALE, ORANGE)

chapter(doc, 6, "端到端模型总览", "把 57 个执行阶段压缩为 6 个职责明确的宏模块。")
add_table(doc, ["宏模块", "输入", "主要工作", "输出"], [
    ["1 数据/地图", "四个原始文件", "异常剔除；LOS/local/rich；信道描述缓存", "有效索引与特征"],
    ["2 矩形 OOF", "位置与几何组", "5 折挖空、官方人口加权", "无泄漏验证协议"],
    ["3 基础谱搜索", "训练信道与特征", "分组 PAS/PDP 邻域；Phase1–6", "Phase6 底座"],
    ["4 锚点与组合", "Phase6 + 岛内锚点", "Phase7–10；残差、门控、组合", "P9 与 P10"],
    ["5 官方反馈", "P9/P10", "冻结 anti-P10 方向", "Phase40"],
    ["6 选择性修正", "Phase40 + g5/g6 锚点", "full-192 PAS + 安全尺度审计", "Phase93"],
], [1.0, 1.45, 3.45, 1.1])
formula(doc, "Raw files → clean/features → rectangular OOF → Phase6 → P9/P10 → Phase40 → Phase93 NPY")

chapter(doc, 7, "地图上下文：点云决定传播相似度", "点云不做完整射线追踪，而是服务于邻域、描述器和门控。")
add_table(doc, ["特征族", "描述内容", "主要用途"], [
    ["LOS", "基站—UE 路径上的遮挡高度、密度、墙面；左右侧分别编码", "地图条件距离、基站侧差异"],
    ["Local", "测试点周围局部高度、墙面/地面密度、多尺度统计", "局部邻域与边界识别"],
    ["Rich", "方向、角度、障碍与频带描述", "树模型、MLP、门控与残差校准"],
], [1.0, 3.65, 2.35])
paragraph(doc, "与 v39 的关系：v39 的 209 维走廊/材质/环形/天际线描述更系统；Round2 没有照搬全部维度，但继承了“地图决定哪些训练点在传播意义上相似”的核心思想，并围绕矩形 OOF 对特征使用方式做更强任务定制。")

chapter(doc, 8, "分组空间预测：PAS 与 PDP 使用不同邻域", "官方几何组不是一个统一平稳空间场。")
paragraph(doc, "推理时在 3738 个有效训练点建立 cKDTree，每个测试点先查询最多 384 个候选。official_island_labels 根据测试位置生成几何组；每个组分别配置 PAS 与 PDP 的邻居数 k、距离幂 p 和广义均值 q。")
formula(doc, "w_j ∝ 1 / max(d_j,0.25)^p,      Σ_j w_j = 1")
bullet(doc, "PAS 与 PDP 可以使用完全不同的 k 和 p，避免角域最优邻居强迫时延域服从。")
bullet(doc, "局部仿射权重使加权邻居坐标中心贴近查询点；部分组使用二次矩约束。")
bullet(doc, "部分组使用角谱圆周矩、二维互相关或 horizontal shift 对齐，避免直接平均错位角谱。")
bullet(doc, "地图条件 metric 与 harmonic graph 允许 XY 较近但遮挡环境不同的点被推远。")
callout(doc, "Round2 的关键变化", "Round1 主要做多专家选择；Round2 进一步把专家适用域显式绑定到官方几何分组，并允许 PAS/PDP 在同一测试点选择不同邻域。", PALE, GREEN)

chapter(doc, 9, "基础复信道：目标谱与交替投影", "局部邻居给出谱形状，复数相位尽量从当前信道保留。")
paragraph(doc, "从 PAS 邻居聚合角域功率目标，从 PDP 邻居聚合时延功率目标。初值可取最近邻、加权复信道或前一阶段通道。一次交替投影先进入 BS 角域替换幅度，再回到信道域进入频率 FFT 替换时延幅度。")
formula(doc, "Z = FFT(H),   Z' = Z · [(1−λ) + λ·sqrt(P_target)/max(|Z|,ε)]")
numbered_list(doc, [
    "BS-FFT：替换角域幅度，使当前信道靠近目标 PAS，保留当前角域复相位。",
    "BS-IFFT：返回天线域。",
    "Frequency-FFT：替换时延域幅度，使当前信道靠近目标 PDP。",
    "Frequency-IFFT：返回复信道，重复有限轮。",
])
paragraph(doc, "PAS 与 PDP 是两个不完全兼容的边缘约束，不能假设无限迭代一定更好。Round2 按几何组冻结投影轮数、初值和最终 PAS blend。")
callout(doc, "幅度规范", "数据存在难以由坐标预测的点级相位规范。基础提交使用保守非零幅度，保持 PAS/PDP 形状并让 NMSE 稳定在接近 1 的区间，避免为不可预测相位过拟合。", WARN, ORANGE)

chapter(doc, 10, "Phase6：共同物理谱底座", "后续 P9 与 P10 都从 Phase6 出发。")
paragraph(doc, "Phase6 的目标是生成更准确的 24-band PAS。24 个频带由 192 子载波每 8 个取一带形成，标签直接从训练信道 BS-FFT 功率得到。")
add_table(doc, ["输入块", "例子", "作用"], [
    ["位置/组", "x、y、side、group", "确定官方几何适用域"],
    ["方向", "horizontal / vertical canonical", "对齐阵列角谱主方向"],
    ["地图", "LOS/local/rich", "区分遮挡与材料上下文"],
    ["信道描述", "24-band PAS/PDP、PCA/统计", "学习局部谱形状"],
    ["模型", "MLP、ExtraTrees、canonical、rich-tree、gate", "矩形 OOF 选型与冻结组合"],
], [1.15, 2.55, 3.3])
callout(doc, "理解方式", "Phase6 决定最终结果的大部分空间结构；后续版本不是重新训练一个全新模型，而是在共同底座上增加有 OOF 证据的可解释方向修正。", PALE, BLUE)

chapter(doc, 11, "P9：锚点残差、收益门控与联合投影", "矩形内部训练点成为官方几何锚点。")
paragraph(doc, "official_anchors 只从非零训练行选择位于测试包围盒内的点，并按几何组隔离。为防止锚点真值直接泄漏，先用矩形外/非锚点训练点预测每个锚点，再计算真值相对基线的残差。")
formula(doc, "r_anchor = clip(log((P_truth+ε)/(P_external+ε)), −2, 2)")
add_table(doc, ["P9 子模块", "参数/做法", "输出"], [
    ["局部 PAS", "horizontal alignment；4 邻居；距离幂 3.0", "local_pas"],
    ["PAS 锚点残差", "岛内最多 16 锚点；24 带残差在频带上取均匀", "pas_residual"],
    ["PDP 锚点残差", "外部 8 邻居基线；岛内最多 4 锚点", "pdp_residual"],
    ["收益门控", "ExtraTrees 500 棵；min_leaf=80；max_features=0.7", "每测试点 alpha"],
    ["联合投影", "PAS residual α=0.15；PDP α=0.025；12 次", "P9 通道"],
], [1.3, 3.85, 1.85])
formula(doc, "PAS_P9 = normalize((1−α)·PAS_corrected + α·PAS_local)")

chapter(doc, 12, "Phase10：构造另一条分组 PAS 方向", "P10 的价值是提供可解释对照，不是最终答案。")
paragraph(doc, "Phase10 从 Phase6 底座构造另一条 PAS 方向：核心组 {1,3,4,9,10} 使用 robust125，补集组 {0,2,5,6,7,8} 使用 graph/canonical/GP 冻结组合；组 4/5/10 再加入 primary-anchor 修正。")
formula(doc, "PAS_target = normalize((1−α)·PAS_base + α·PAS_component)")
add_table(doc, ["部分", "冻结设置"], [
    ["核心组", "α=1.25，robust_binary_all 子集"],
    ["补集组", "α=1.0，graph_canonical_gp 冻结组合"],
    ["主锚点", "组 4/5/10；residual α=0.10；local scale=0.50；clip=0.30"],
    ["投影", "24 bands；4 次；PAS ratio clip=[0.25,4.0]；PDP 固定为 Phase6"],
], [1.45, 5.55])

chapter(doc, 13, "Phase40：把官方反馈冻结为 anti-P10", "官方测试反馈改变了优化方向。")
paragraph(doc, "已知官方结果中 P9 约 0.6395，P10 约 0.6354。P10 方向在官方隐藏测试上持续变差，因此 Phase40 不向 P10 靠近，而是把 P9→P10 的对数谱方向取反。")
formula(doc, "d = clip(log((PAS_P10+ε)/(PAS_P9+ε)), −2, 2)")
formula(doc, "PAS_Phase40 = normalize(PAS_P9 · exp(−0.50·d))")
paragraph(doc, "24-band 目标重复到 192 个子载波，在 BS-FFT 域按 sqrt(P_target/P_current) 修正幅度，再 IFFT 回到复信道。eta=0.50 是冻结参数，不因后续分数再搜索。")
callout(doc, "重要口径", "Phase40 是“官方反馈黑盒方向”的可复现编码，而不是本地 OOF 最优模型。它解释了为什么后续优化不能继续单纯追逐更高本地分数。", WARN, RED)

chapter(doc, 14, "Phase93：只修正 g5/g6 的 89 个测试点", "有证据才扩大修正，其他 411 点锁定 Phase40。")
paragraph(doc, "Phase93 先从原始 P9 复制通道，只对几何组 g5 和 g6 增加 full-192 PAS 锚点目标。两组共有 43+46=89 个测试点，训练锚点分别为 22 和 42 个。")
formula(doc, "PAS_192_desired = normalize(0.80·PAS_192_P9 + 0.20·PAS_192_anchor)")
numbered_list(doc, [
    "对锚点 full-192 PAS 做 horizontal alignment。",
    "每个测试点使用 4 个同组锚点，距离平方权重插值。",
    "对 89 行执行 12 次 full-192 角域投影。",
    "仍使用 P9/P10 计算的冻结 anti-P10 方向。",
    "其余 411 行必须与 Phase40 bitwise 相同。",
])
add_table(doc, ["冻结参数", "值"], [
    ["active groups", "[5,6]"], ["active test rows", "89"], ["anchor counts", "g5=22，g6=42"],
    ["neighbors / power", "4 / 2.0"], ["dose / iterations", "0.20 / 12"], ["anti eta", "0.50"],
], [2.4, 4.6])
callout(doc, "设计思想", "提交次数有限时，优先扩大矩形 OOF 中有稳定证据的局部修正，不把风险扩散到全部 500 点。", PALE, GREEN)

chapter(doc, 15, "symmetric clamp-floor：正式输出中实际未生效", "必须把“设计机制”和“最终事实”分开讲。")
paragraph(doc, "clamp-floor 只从原始 P9 预测计算每个测试样本×UE 分支的 PAS/PDP 低尾范数，不读取测试真值。")
formula(doc, "q = min(q01(PAS_norm), q01(PDP_norm))")
formula(doc, "s = clip(sqrt(1e−15 / max(q,1e−30)), 1, 5)")
paragraph(doc, "若 s>1，则把该 UE 分支在全部 256 天线和 192 子载波上等比例抬升；若 s=1，必须 bitwise 不变。正式 Phase93 manifest 中 2000 个 UE 分支全部 s=1，active_scale_branches=0，因此 clamp-floor 没有改变任何正式测试数值。")
callout(doc, "最终事实", "正式 Phase93 输出等于 Phase89 数值：g5/g6 full-192 修正后再应用冻结 anti-P10。clamp-floor 只是安全兜底，不是得分提升来源。", WARN, RED)

chapter(doc, 16, "训练、冻结与防泄漏", "复杂工程必须有简单、可审计的证据链。")
bullet(doc, "所有模型选择以 5 折矩形 OOF 为依据，并按官方几何组人口加权。")
bullet(doc, "锚点真值先与矩形外预测比较形成残差，再插值到测试点；测试真值从未出现。")
bullet(doc, "Phase93 在读取 Phase92/Phase93 交互分数前预声明 dose、组、顺序和通过条件。")
bullet(doc, "构建器拒绝覆盖已有输出，先写 .building.npy，再原子替换。")
bullet(doc, "manifest 固定输入脚本哈希、中间数组哈希、输出 SHA256 和 bitwise 不变量。")
add_table(doc, ["完整性约束", "要求"], [
    ["P9/P40/P89/P90 复现", "最大指标误差分别 ≤1e−10 / 1e−10 / 1e−10 / 1e−8"],
    ["非 g5/g6 行", "411 行与 Phase40 bitwise 相同"],
    ["s=1 UE 分支", "与 Phase89 bitwise 相同"],
    ["输出", "finite、500 行非零、shape/dtype/bytes/SHA256 固定"],
], [2.05, 4.95])

chapter(doc, 17, "以一个测试点走完整流程", "把所有模块串成一个可复述的端到端故事。")
numbered_list(doc, [
    "读取测试坐标 q，从 test_pos 的 DBSCAN 标签确定几何组，并查询 LOS/local/rich 地图特征。",
    "在有效训练点中查询候选邻居；按该组 PAS/PDP 的 k、p、q 与 metric 分别构造两个权重计划。",
    "聚合邻居 PAS/PDP，做角谱对齐、仿射/二次矩修正或 harmonic graph，交替投影形成基础复信道。",
    "Phase6 根据位置、方向、地图和 24-band 描述生成共同 PAS 底座。",
    "P9 使用同组官方锚点残差与 ExtraTrees 收益门控，12 次联合投影得到 P9。",
    "Phase10 从 Phase6 构造分组组合 PAS，用作官方反馈方向的另一端点。",
    "Phase40 计算 P9→P10 的对数方向并反向移动 eta=0.50。",
    "若测试点属于 g5/g6，则用 4 个锚点生成 full-192 PAS，dose=0.20 投影 12 次，再应用 anti-P10；否则锁定 Phase40。",
    "从 P9 计算 clamp-floor；正式测试中 s=1，因此不改变通道。",
    "写入 complex64 输出并检查 shape、finite、非零、文件大小和 SHA256。",
])
callout(doc, "一句话复述", "几何组决定“当前点属于哪套空间规律”，PAS/PDP 邻域决定“目标谱长什么样”，交替投影决定“怎样把谱写回复信道”，锚点与 anti-P10 决定“局部应向哪个方向修正”。", PALE, BLUE)

chapter(doc, 18, "结果、复现与提交", "只陈述可复现证据，不把本地分数当官方分数。")
add_table(doc, ["项目", "结果/要求"], [
    ["Phase93 矩形 OOF 相对 Phase40", "平均 +0.000806；5 折中 4 折为正；最差 −0.000130"],
    ["稳健下界", "mean − 0.75×std = +0.000272"],
    ["已知官方参照", "Phase40 官方测试 0.639697；仓库未记录 Phase93 官方隐藏测试分"],
    ["输出 shape / dtype", "(500,256,4,192) / complex64"],
    ["文件大小", "786,432,128 bytes"],
    ["正式 SHA256", "047b821a5dc6b02abb7fe99c899b2060e52f2b14504a2c2e1605b84adb4201e8"],
], [2.3, 4.7])
doc.add_heading("18.1 从原始数据运行", level=2)
formula(doc, "python -m pip install -r requirements.txt")
formula(doc, "python build_phase93_end_to_end.py")
paragraph(doc, "可用 --list-stages 查看 57 个阶段；中断后可用 --from-stage <name> 复用此前缓存继续。最终输出：Round2_Test_Channel_phase93_g56_antip10_plus_symmetric_clamp.npy。")
doc.add_heading("18.2 提交内容", level=2)
bullet(doc, "完整源码与 requirements.txt。")
bullet(doc, "冻结配置、validation JSON、submission manifest 和 raw-to-final 入口。")
bullet(doc, "最终 Phase93 NPY。")
bullet(doc, "赛事原始训练信道、训练/测试位置和点云不重复上传。")

doc.add_heading("附录 A：代码阅读地图", level=1)
add_table(doc, ["顺序", "文件", "关注点"], [
    ["1", "build_phase93_from_raw.py", "57 阶段编排、原始输入验证、最终 QA"],
    ["2", "r2_pipeline.py", "分组邻域、矩形 split、谱投影、地图 metric"],
    ["3", "build_phase6_submission.py", "共同物理 PAS 底座"],
    ["4", "build_phase9_submission.py", "官方锚点、残差、ExtraTrees gate、联合投影"],
    ["5", "build_phase10_robust125_anchor_pas_submission.py", "分组 PAS 组合与 Phase10 通道"],
    ["6", "build_phase40_p9_actual_antip10_pas050_submission.py", "anti-P10 官方反馈方向"],
    ["7", "build_phase93_g56_antip10_plus_symmetric_clamp_submission.py", "g5/g6 full-192、clamp、bitwise QA"],
], [0.55, 3.1, 3.35])

doc.add_heading("附录 B：常见答辩问题", level=1)
add_table(doc, ["问题", "回答"], [
    ["这与 Round1 v39 是否相同？", "共享评分对齐谱与复信道交替投影主线；Round2 新增矩形 OOF、几何分组、锚点残差、官方反馈方向和 89 点选择性修正。"],
    ["为什么不用随机验证？", "随机验证点周围仍有密集训练邻居，会高估矩形空洞上的空间插值能力。"],
    ["为什么不直接训练大网络？", "样本只有 3738 个有效点，而每点有 196,608 个复数；谱表征与局部几何更贴合指标并更可审计。"],
    ["Phase93 得分提升来自 clamp-floor 吗？", "不是。正式测试 2000 个 UE 分支全部 s=1；最终改变来自 g5/g6 full-192 与 anti-P10。"],
    ["0.639697 是 Phase93 官方分吗？", "不是，是 Phase40 的已知官方参照。仓库未记录 Phase93 官方隐藏测试分。"],
    ["赛事方能否从原始数据重建？", "可以。四个官方原始文件放入仓库根目录，运行 build_phase93_end_to_end.py。"],
], [2.1, 4.9])

callout(doc, "最终总结", "Phase93 的价值不在版本数量，而在把 Round1 的谱重建主线改造成适配 Round2 官方矩形几何、具备无泄漏证据链和端到端复现能力的完整系统。", PALE, GREEN)

REPO_OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(REPO_OUT)
doc.save(DESKTOP_OUT)
print(REPO_OUT)
print(DESKTOP_OUT)
