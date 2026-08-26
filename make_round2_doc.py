from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\asus\Desktop\Round2_完整版.docx")
BLUE = "2E74B5"
DARK = "0B2545"
ORANGE = "D96C32"
LIGHT = "F2F5F8"
CALLOUT = "E8EEF5"
WARN = "FFF2CC"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def mark_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    flag = OxmlElement("w:tblHeader")
    flag.set(qn("w:val"), "true")
    tr_pr.append(flag)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    flag = OxmlElement("w:cantSplit")
    flag.set(qn("w:val"), "true")
    tr_pr.append(flag)


def put(cell, value, bold=False, color=None, size=9.0):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1.5)
    r = p.add_run(str(value))
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, heads, rows, widths):
    t = doc.add_table(rows=1, cols=len(heads))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    mark_header(t.rows[0])
    for i, head in enumerate(heads):
        put(t.rows[0].cells[i], head, True, DARK, 9.0)
        shade(t.rows[0].cells[i], LIGHT)
    for values in rows:
        cells = t.add_row().cells
        for i, value in enumerate(values):
            put(cells[i], value)
    for row in t.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def step(doc, text):
    doc.add_paragraph(text, style="List Number")


def numbered_step(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    lead = p.add_run(f"{number}.  ")
    lead.bold = True
    p.add_run(text)
    return p


def formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Cambria Math"
    r.font.size = Pt(11.5)
    r.font.color.rgb = RGBColor.from_string(DARK)


def callout(doc, label, text, fill=CALLOUT):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    prevent_row_split(t.rows[0])
    c = t.cell(0, 0)
    shade(c, fill)
    c.text = ""
    p = c.paragraphs[0]
    a = p.add_run(label + "：")
    a.bold = True
    a.font.color.rgb = RGBColor.from_string(ORANGE)
    b = p.add_run(text)
    b.font.color.rgb = RGBColor.from_string(DARK)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.78)
sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(0.85)
sec.right_margin = Inches(0.85)
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12
for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11, DARK, 8, 4),
]:
    style = doc.styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(28)
r = p.add_run("Round2 Phase93 端到端模型算法说明")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = RGBColor.from_string(DARK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("从赛事原始训练数据到 500 点复数信道输出")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor.from_string(BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("面向赛事复核、代码审计与答辩汇报")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor.from_string("666666")
callout(
    doc,
    "严格口径",
    "当前 build_phase93_end_to_end.py 已改为严格 raw-to-final 入口：只读取赛事提供的训练信道、训练位置、测试位置和地图点云。train_energy、地图特征、矩形 OOF 折、PAS/PDP 标签、Phase4/5/6、P9、P10、Phase40 与 Phase93 均由代码现场生成；历史 Release 数组仅用于结果审计，不参与建模或推理。",
    CALLOUT,
)
doc.add_page_break()

doc.add_heading("0. 先理解 PAS、PDP 与 NMSE", 1)
doc.add_paragraph(
    "Phase93 的输入和输出都是复数信道 H。H[a,u,f] 表示第 a 个基站天线端口、第 u 个 UE 天线/极化分支、"
    "第 f 个子载波上的复数响应。复数的模表示幅度，辐角表示相位。赛事指标不只比较每个复数元素，"
    "还分别比较信道在阵列角度域和时延域中的功率分布。"
)
table(doc, ["指标", "回答的问题", "变换维度", "越大/越小越好"], [
    ["PAS", "能量主要从哪些阵列角度到达？", "256 个 BS 端口重排为 2×16×8，沿 16×8 做 FFT2", "PAS 余弦相似度越接近 1 越好"],
    ["PDP", "能量主要落在哪些传播时延上？", "沿 192 个子载波做 FFT", "PDP 余弦相似度越接近 1 越好"],
    ["NMSE", "预测复信道与真值逐元素相差多少？", "直接比较复数 H_pred 与 H_true", "NMSE 越接近 0 越好"],
], [0.75, 2.05, 2.75, 1.35])

doc.add_heading("0.1 PAS：阵列角度功率谱", 2)
doc.add_paragraph(
    "PAS（Power Angular Spectrum，阵列角度功率谱）描述不同空间角度上的接收功率。"
    "本题的 256 个基站端口按 2 个极化、16 个水平阵元、8 个垂直阵元重排。"
    "对 16×8 阵列维做二维正交 FFT，相当于把“天线位置域”转换到“空间波束/角度域”。"
)
formula(doc, "A = FFT2_BS(H),    P_pas = |A|^2")
doc.add_paragraph(
    "代码对每个样本、每个 UE 分支和每个子载波，沿 256 个角度 bin 计算归一化功率方向。"
    "预测 PAS 与真值 PAS 的相似度为："
)
formula(doc, "C_pas = mean[(P_pred · P_true) / (||P_pred||_2 ||P_true||_2)]")
callout(
    doc,
    "PAS 小例子",
    "用 4 阵元的一维简化阵列说明：若 H=[1,1,1,1]，空间 FFT 的能量集中在第 0 个角度 bin，"
    "归一化 PAS 可写成 [1,0,0,0]。若预测也为 [1,1,1,1]，两个 PAS 的余弦相似度为 1；"
    "若预测 H=[1,-1,1,-1]，能量移到另一个角度 bin，PAS 近似 [0,0,1,0]，与真值的余弦相似度为 0。"
)
doc.add_paragraph(
    "PAS 余弦相似度只关心功率谱形状：即使预测 PAS 是真值的 3 倍，归一化后的方向仍相同，C_pas 仍为 1。"
    "因此 PAS 擅长评价角度结构，但不能单独约束总幅度。"
)

doc.add_heading("0.2 PDP：频率/时延功率谱", 2)
doc.add_paragraph(
    "PDP（Power Delay Profile，功率时延谱）描述多径能量分布在不同传播时延上的情况。"
    "192 个子载波处在频率域；沿子载波维执行正交 FFT 后，频率变化被转换成离散时延 bin。"
)
formula(doc, "D = FFT_frequency(H),    P_pdp = |D|^2")
doc.add_paragraph(
    "代码对每个样本、每个 BS 端口和每个 UE 分支，沿 192 个时延 bin 计算预测与真值的余弦相似度，再取平均："
)
formula(doc, "C_pdp = mean[(D_pred_power · D_true_power) / (||D_pred_power||_2 ||D_true_power||_2)]")
callout(
    doc,
    "PDP 小例子",
    "用 4 个子载波简化说明：若频域信道 H=[1,1,1,1]，FFT 后能量集中在第 0 个时延 bin，"
    "归一化 PDP 为 [1,0,0,0]，表示主要是零时延/共同相位分量。若 H=[1,-1,1,-1]，"
    "能量集中到另一个时延 bin，PDP 位置发生移动。预测与真值集中在同一 bin 时 C_pdp=1；"
    "集中在不同 bin 时相似度接近 0。"
)
doc.add_paragraph(
    "PAS 与 PDP 从两个不同方向约束同一个复信道：PAS 检查阵列空间结构，PDP 检查子载波随频率变化形成的时延结构。"
    "这也是模型需要交替投影的原因——只匹配其中一个谱，可能破坏另一个谱。"
)

doc.add_heading("0.3 NMSE：原始复信道归一化均方误差", 2)
doc.add_paragraph(
    "NMSE（Normalized Mean Squared Error，归一化均方误差）直接比较预测复信道与真值复信道。"
    "复数误差的平方模同时包含实部和虚部差异：|x+jy|^2=x^2+y^2。"
)
formula(doc, "NMSE = sum |H_pred - H_true|^2 / sum |H_true|^2")
callout(
    doc,
    "NMSE 小例子",
    "设真值 H_true=[1+j, 1]，预测 H_pred=[0.8+1.1j, 0.9]。误差能量为 "
    "|-0.2+0.1j|^2+|-0.1|^2=0.05+0.01=0.06；真值能量为 |1+j|^2+|1|^2=2+1=3；"
    "所以 NMSE=0.06/3=0.02。对应评分中的 NMSE 项为 0.2/(1+0.02)=0.1961，接近该项上限 0.2。"
)
doc.add_paragraph(
    "NMSE=0 表示复数信道逐元素完全一致。NMSE 不仅受功率谱形状影响，也对整体幅度和相位敏感。"
    "例如 H_pred=-H_true 时，PAS 与 PDP 完全相同，但复数误差很大；因此三项指标需要共同使用。"
)

doc.add_heading("0.4 三项指标如何形成总分", 2)
formula(doc, "score = 0.4*C_pas + 0.4*C_pdp + 0.2/(1+NMSE)")
doc.add_paragraph(
    "例如 C_pas=0.80、C_pdp=0.70、NMSE=1.00，则总分为 0.4×0.80 + 0.4×0.70 + 0.2/(1+1) "
    "= 0.32 + 0.28 + 0.10 = 0.70。PAS 和 PDP 各占 40%，所以 Phase93 的主体优化集中在角度谱与时延谱；"
    "NMSE 占 20%，用于约束预测信道不要在复数幅度和相位上偏离真值过远。"
)
callout(
    doc,
    "阅读后续章节时的抓手",
    "邻域、树模型和锚点负责预测 PAS/PDP 目标；交替 FFT/IFFT 投影负责把这两个目标写回复数信道；"
    "最终幅度策略与 clamp-floor 主要控制 NMSE 风险和异常低能量分支。"
)

doc.add_heading("1. 任务定义与输入输出契约", 1)
doc.add_paragraph("目标是利用 4000 个带标签训练采样点，预测 500 个测试采样点的完整宽带多天线复数信道。每个样本包含 256 个基站阵列端口、4 个 UE 天线/极化分支和 192 个子载波。")
table(doc, ["对象", "典型形状/类型", "来源", "处理规则"], [
    ["Round2_Train_Channel.npy", "(4000,256,4,192), complex", "赛事原始数据", "唯一信道监督；按块 mmap 读取"],
    ["Round2_Train_Pos.npy", "(4000,3), float", "赛事原始数据", "平面坐标建邻域；三维坐标用于方向特征"],
    ["Round2_Test_Pos.npy", "(500,3), float", "赛事原始数据", "DBSCAN 识别测试矩形岛；无信道真值"],
    ["Round2_Map.ply", "ASCII PLY，三维场景网格", "赛事原始数据", "生成 LOS、遮挡和局部环境特征"],
    ["train_energy.npy", "(4000,), float64", "由训练信道生成", "缓存，不是额外监督"],
    ["最终输出", "(500,256,4,192), complex64", "模型生成", "所有值有限、所有测试行非零"],
], [1.45, 1.55, 1.25, 2.55])

doc.add_heading("2. 原始信道清洗与异常点剔除", 1)
doc.add_paragraph("对每个训练点计算整张量能量：")
formula(doc, "E_i = sum |H_i[a,u,f]|^2,   a=1..256, u=1..4, f=1..192")
doc.add_paragraph("若 E_i=0，则整点信道响应为全零。Phase93 代码将其视为 outlier，不参与邻域索引、模型训练、锚点构造或验证。历史数据中共有 262 个零能量点，最终有效训练点为 3738 个。train_energy.npy 只是该计算的缓存；严格重建时可从 Round2_Train_Channel.npy 自动生成。")

doc.add_heading("3. 非随机矩形验证集", 1)
doc.add_paragraph("测试点不是训练分布的随机子样本，而是集中在多个小矩形中。若随机抽点验证，验证点附近仍有大量训练点，会高估空间插值能力。因此验证必须制造与官方测试相似的空间空洞。")
for number, text_value in enumerate([
    "使用测试位置在二维平面执行 DBSCAN（eps=10、min_samples=3），恢复官方测试岛标签和每个岛的包围矩形。",
    "在训练点两侧分别选择矩形候选中心；宽、高和目标人口参考 TEST_BLOCKS，并乘 scale=0.75。",
    "候选矩形内必须是非零信道点，并优先使实际人口接近目标人口；矩形之间不重叠。",
    "每个验证矩形外再加 3 m buffer，从训练邻域中制造真实的空间缺口。",
    "用固定种子 20260813+fold 构造 5 折；按官方各岛人口对折内样本加权。",
], start=1):
    numbered_step(doc, number, text_value)
callout(doc, "目的", "验证的重点是模拟官方外推几何，而不是追求普通随机交叉验证的低方差。")

doc.add_heading("4. 评估指标与建模空间", 1)
doc.add_paragraph("模型同时优化三类性质：阵列角度功率谱 PAS、频率/时延功率谱 PDP，以及原始复信道 NMSE。")
formula(doc, "score = 0.4*C_pas + 0.4*C_pdp + 0.2/(1+NMSE)")
doc.add_paragraph("C_pas 和 C_pdp 均为归一化功率谱余弦相似度。由于复信道的点相位规范难以仅由坐标预测，模型把主要容量放在 PAS/PDP 形状上，并用保守的幅度策略避免 NMSE 项失控。")

doc.add_heading("5. 从原始信道构造频谱监督", 1)
doc.add_heading("5.1 基站二维阵列 PAS", 2)
doc.add_paragraph("将 256 个基站端口重排为 2 个极化 × 16 个水平阵元 × 8 个垂直阵元，对 16×8 阵列维执行正交二维 FFT：")
formula(doc, "A = FFT2_BS(H);   PAS_192 = |A|^2 / |||A|^2||_antenna")
doc.add_paragraph("将 192 个子载波分成 24 个频带，每带 8 个子载波取均值，再沿 256 阵列维归一化，得到 (N,256,4,24) 的 PAS_band24。")
doc.add_heading("5.2 子载波 FFT 与 PDP", 2)
formula(doc, "D = FFT_frequency(H);   PDP_192 = |D|^2 / |||D|^2||_delay")
doc.add_paragraph("PAS 描述阵列角度能量，PDP 描述时延能量。后续投影不会直接回归每个复数元素，而是在这两个功率约束之间交替修正复数信道。")

doc.add_heading("6. 原始数据基础预测：分岛邻域 + 交替频谱投影", 1)
doc.add_paragraph("r2_pipeline.generate_submission() 给出了从原始训练信道到基础测试信道的核心机制。首先根据测试矩形岛分别建模；PAS 与 PDP 可以选择不同的邻居集合、邻居数和距离指数。")
for text_value in [
    "在 3738 个有效训练点上建立 cKDTree；每个测试点最多查询 384 个候选邻居。",
    "按测试岛选择 PAS/PDP 的 k、距离幂次和投影混合参数；基础权重为 w_j proportional to 1/max(d_j,0.25)^p。",
    "用局部仿射/二次修正使加权坐标中心更贴近查询点；部分岛使用特征距离或 harmonic graph 邻域。",
    "从邻居训练信道聚合目标 PAS 与目标 PDP；初值取最近邻或复信道加权和。",
    "在 BS-FFT 域匹配 PAS 幅度，在频率 FFT 域匹配 PDP 幅度，交替 IFFT 返回复数域。",
    "最后用邻居能量估计统一幅度；采用保守小幅度时，PAS/PDP 形状不变而 NMSE 接近稳定区间。",
]:
    bullet(doc, text_value)
callout(doc, "基础模型本质", "空间邻域决定“应该像谁”，频谱投影决定“怎样变成同时满足角度和时延约束的复数张量”。")

doc.add_heading("7. Phase6：物理频谱描述增强", 1)
doc.add_paragraph("Phase6 在基础复信道上学习更准确的 24-band PAS 目标。训练标签直接来自第 5 节的原始训练信道频谱缓存；特征来自位置、两侧基站相对方向、LOS/rich-map 信息和局部结构。")
table(doc, ["组件", "算法逻辑", "冻结参数"], [
    ["方向标准化", "按左右场景使用基站坐标，将位置转为单位视线方向；估计水平/垂直循环移位并对齐 PAS", "水平周期16，垂直周期8"],
    ["低维回归", "对 sqrt(PAS) 展平后 PCA；ExtraTrees 从几何特征回归 PCA 系数，再逆变换与平方", "PCA=160；ET约420-500棵"],
    ["物理组合", "融合基础、水平拟合、垂直拟合、安全组合与 rich-tree 目标，并做非负归一化", "24 bands"],
    ["复信道投影", "目标 PAS 比例限制后，交替保持基础 PDP 与目标 PAS", "scale=1.25；PDP power=1.5；12次"],
], [1.1, 4.25, 1.45])
doc.add_paragraph("Phase6 输出为 Round2_Test_Channel_matched_phase6_delta2053.npy。它是 P9 和 Phase10 的共同底座，决定最终结果的大部分空间结构。")

doc.add_heading("8. P9：锚点残差、局部 PAS 与门控联合投影", 1)
doc.add_paragraph("P9 利用测试矩形内存在的训练点作为 official anchors。锚点只来自非零训练行；对每个测试岛单独计算，避免不同矩形之间串扰。")
for text_value in [
    "基础 PAS：从 Phase6 复信道计算 24-band PAS。局部 PAS：用水平对齐后的 4 邻居、距离幂 3.0 进行局部预测。",
    "PAS 锚点残差：log((PAS_truth+eps)/(PAS_external+eps)) 裁剪到 [-2,2]，对 24 带求均匀残差，再在岛内最多 16 个锚点插值。",
    "PDP 锚点残差：用岛外/非锚点训练样本建立 8 邻居 PDP 基线，再取 log 残差并在岛内最多 4 个锚点插值。",
    "门控模型：ExtraTreesRegressor，500 棵树，min_samples_leaf=80，max_features=0.7，random_state=52180；根据 5 折收益网格选择 alpha。",
    "alpha=clip(0.75*grid[argmax(gate)],0,0.6)；PAS residual alpha=0.15，PDP residual alpha=0.025。",
    "从 Phase6 复信道出发执行 12 次 PAS/PDP 联合交替投影，得到 P9。",
]:
    bullet(doc, text_value)

doc.add_heading("9. Phase10：分组组合 PAS 目标", 1)
doc.add_paragraph("Phase10 不是直接沿用 P9，而是从 Phase6 底座构造另一条 PAS 方向，作为后续 anti-P10 的参照。核心组 {1,3,4,9,10} 使用 robust125，其他组使用 graph/canonical/GP 冻结组合；组 4/5/10 再加入 primary-anchor 修正。")
formula(doc, "PAS_target = normalize((1-alpha)*PAS_base + alpha*PAS_component)")
doc.add_paragraph("核心组 alpha=1.25，补集组 alpha=1.0；primary-anchor residual alpha=0.10，local scale=0.50，local clip=0.30。生成通道时以 Phase6 为初值，做 4 次交替投影，PAS ratio clip=[0.25,4.0]，PDP 约束始终取 Phase6 PDP。")

doc.add_heading("10. Phase40：用官方反馈冻结 anti-P10 方向", 1)
doc.add_paragraph("将 P9 与 P10 都变换为归一化 24-band PAS，定义两者的对数方向：")
formula(doc, "d = clip(log((PAS_P10+eps)/(PAS_P9+eps)), -2, 2)")
doc.add_paragraph("Phase40 不朝 P10 靠近，而是沿相反方向移动：")
formula(doc, "PAS_desired = normalize(PAS_current * exp(-0.50*d))")
doc.add_paragraph("将 24 带目标重复到 192 个子载波，在 BS-FFT 域按 sqrt(P_target/P_current) 修正幅度，再 IFFT 回复数信道。eta=0.50 是冻结的官方反馈探针参数。")

doc.add_heading("11. Phase93：g5/g6 full-192 局部修正", 1)
doc.add_paragraph("Phase93 只对测试岛 5 和 6（共 89 个测试点）增加 full-192 PAS 锚点目标。训练锚点分别为 22 和 42 个；先做 horizontal alignment，再以 4 邻居、距离平方权重插值。")
formula(doc, "PAS_192_desired = normalize(0.80*PAS_192_P9 + 0.20*PAS_192_anchor)")
doc.add_paragraph("对这 89 行执行 12 次 full-192 投影；随后仍使用第 10 节冻结的 anti-P10 方向。其余 411 行必须与 Phase40 bitwise 一致，以限制局部修正的外溢。")

doc.add_heading("12. symmetric clamp-floor 与最终事实", 1)
doc.add_paragraph("最后仅从原始 P9 预测计算每个测试样本、每个 UE 分支的 PAS/PDP 范数低尾，不读取任何测试真值。")
formula(doc, "q=min(q01(PAS_norm),q01(PDP_norm));  s=clip(sqrt(1e-15/max(q,1e-30)),1,5)")
doc.add_paragraph("若 s>1，则把该 UE 分支在全部 256 天线和 192 子载波上等比例抬升；若 s=1，则保持 bitwise 不变。最终 Phase93 manifest 显示 2000 个 UE 分支全部 s=1，因此这次正式输出中 clamp-floor 实际没有改变任何数值；最终输出等于完成 g5/g6+anti-P10 后的 Phase89 结果。")

doc.add_heading("13. 完整执行伪代码", 1)
for text_value in [
    "读取 train_channel/train_pos/test_pos；计算 train_energy；valid=energy>0。",
    "从 test_pos 聚类得到岛标签和矩形；构造 5 个矩形验证折并按官方人口加权。",
    "从训练信道计算 PAS_192、PAS_band24、PDP_192 标签缓存。",
    "按岛建立 PAS/PDP 邻域，交替频谱投影得到基础复信道。",
    "训练 Phase6 的 canonical/vertical/PCA-ExtraTrees PAS 目标并投影回复信道。",
    "用锚点残差、局部预测和 ExtraTrees gate 生成 P9 PAS/PDP 目标；12 次联合投影得到 P9。",
    "构造 Phase10 分组 PAS 组合；以 Phase6 PDP 为约束做 4 次投影得到 P10。",
    "由 P9/P10 计算 anti-P10 方向，eta=0.50 得到 Phase40。",
    "仅对 g5/g6 生成 full-192 anchor 目标，dose=0.20，12 次投影；再应用 anti-P10。",
    "计算 P9-only clamp-floor；写出 complex64；检查 shape、finite、nonzero、SHA256。",
]:
    step(doc, text_value)

doc.add_heading("14. 严格端到端复现边界", 1)
table(doc, ["类别", "实际内容", "是否作为模型输入", "说明"], [
    ["赛事原始输入", "Train_Channel、Train_Pos、Test_Pos、Map.ply", "是", "四个文件放在仓库根目录"],
    ["现场生成缓存", "energy、map features、矩形折、PAS/PDP、各阶段 OOF/test 数组", "否", "均可删除后重算；不是提交前置资产"],
    ["固定模型定义", "源码、随机种子、超参数、分组规则和选择配置", "是", "属于可审计的算法与模型参数"],
    ["历史 Release 输出", "Phase6/P9/P10/Phase40/Phase93 参考 npy", "否", "只用于 SHA/数值对照，入口不会下载或读取"],
], [1.3, 2.45, 1.35, 1.95])
callout(doc, "一键入口", "运行 python build_phase93_end_to_end.py。入口顺序执行 57 个确定性阶段，并支持 --list-stages 查看完整计划、--from-stage <name> 在中断后复用已有缓存继续。最终写出 500×256×4×192 的 complex64 Phase93 文件及 raw_end_to_end manifest。", CALLOUT)

doc.add_heading("15. 输出验收与提交内容", 1)
for text_value in [
    "最终文件 shape=(500,256,4,192)，dtype=complex64，文件大小 786,432,128 字节。",
    "所有元素 finite；每个测试行能量大于 0；禁止覆盖已有输出，先写 .building 再原子替换。",
    "正式 Phase93 SHA256：047b821a5dc6b02abb7fe99c899b2060e52f2b14504a2c2e1605b84adb4201e8。",
    "向赛事方提交源码、requirements、冻结配置/manifest、raw-to-final 入口；赛事原始训练/测试数据不重复上传。",
    "若赛事方只要求评分文件，提交 Round2_Test_Channel_phase93_g56_antip10_plus_symmetric_clamp.npy。",
]:
    bullet(doc, text_value)

doc.add_heading("附录：关键代码定位", 1)
table(doc, ["源码", "职责"], [
    ["r2_pipeline.py", "原始数据读取、能量、矩形验证、指标、分岛邻域与频谱投影基础算子"],
    ["build_phase93_from_raw.py", "四个原始文件到 Phase93 的 57 阶段编排、恢复运行与最终验收"],
    ["matched_phase5_tree_band_descriptor.py", "从原始训练信道生成 24-band PAS 缓存"],
    ["matched_phase6_pas_fitted_canonical.py / vertical.py", "方向对齐、PCA+ExtraTrees PAS 回归"],
    ["matched_phase6_physics_combo_channel_validation.py", "Phase6 物理 PAS/PDP 交替投影与冻结参数验证"],
    ["build_phase9_submission.py", "锚点 PAS/PDP 残差、门控、12 次联合投影"],
    ["phase10_pas_complement_anchor_confirmation.py", "Phase10 核心/补集/primary-anchor PAS 目标构造"],
    ["build_phase10_robust125_anchor_pas_submission.py", "Phase10 复信道生成与审计"],
    ["build_phase40_p9_actual_antip10_pas050_submission.py", "Phase40 anti-P10 投影"],
    ["build_phase93_g56_antip10_plus_symmetric_clamp_submission.py", "g5/g6、anti-P10、clamp-floor 与最终 QA"],
], [3.25, 3.55])

for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = footer.add_run("Round2 Phase93 | Raw-to-output algorithm description")
    rr.font.size = Pt(8)
    rr.font.color.rgb = RGBColor.from_string("777777")
doc.core_properties.title = "Round2 Phase93 端到端模型算法说明"
doc.core_properties.subject = "Raw competition data to Phase93 complex channel output"
doc.core_properties.comments = "Strict raw-to-final Phase93 reconstruction; historical prediction arrays are audit-only."
doc.save(OUT)
print(OUT)
